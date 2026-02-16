#!/usr/bin/env python3
"""Generate CheckIn — Product Overview.docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Style setup ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)

def add_bold_paragraph(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table


# ═══════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════

# ── Title ─────────────────────────────────────────────────────
title = doc.add_heading('CheckIn — Product Overview', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

meta = doc.add_paragraph()
meta.add_run('Last updated: February 16, 2026').italic = True
meta.add_run('\nStatus: ').italic = True
meta.add_run('CANONICAL').bold = True
meta.add_run(' — all other documents reference this file for product definitions').italic = True
meta.add_run('\nSources: ').italic = True
meta.add_run('product/overview.md (Git repo), Check-in Thesis V1, Checking V2-Working Doc').italic = True

doc.add_paragraph('─' * 60)

# ── One-Liner ─────────────────────────────────────────────────
doc.add_heading('One-Liner', level=2)
p = doc.add_paragraph()
run = p.add_run('CheckIn is an AI-native customer engagement platform that helps F&B venues own their customer relationships through gamified experiences, a dual credit economy, and per-venue AI agents.')
run.bold = True

# ── The Problem ───────────────────────────────────────────────
doc.add_heading('The Problem', level=2)

doc.add_paragraph(
    "India's F&B industry is bleeding margin in two directions — and losing its "
    "customer relationships in the process."
)

add_bold_paragraph('The Delivery Tax')
doc.add_paragraph(
    "Venues pay 15-28% commission to Zomato and Swiggy on every delivery order, plus payment "
    "gateway fees, GST, platform fees (\u20b910-12/order), and packaging costs. Total effective cost "
    "can exceed 35% of order value. In exchange, the aggregator provides genuine logistics value — "
    "riders, routing, customer support. But the customer's loyalty stays with the app, not the venue. "
    "If a venue leaves Zomato, its delivery customers don't follow."
)

add_bold_paragraph('The Dine-In Trap')
doc.add_paragraph(
    "Zomato Pay and Swiggy Dineout require participating venues to offer 15-40% mandatory discounts "
    "to customers, plus pay 4-12% commission per transaction. The critical issue: these discounts apply "
    "to any customer who pays through the app — including walk-ins who discovered the venue on their own, "
    "not through the aggregator. A customer sits down at your caf\u00e9 because their friend recommended it, "
    "opens Zomato Pay, and the venue pays a 20% discount plus commission to an aggregator that contributed "
    'nothing to that visit. As the NRAI put it: "Why should a restaurant pay a commission to a middleman '
    'to offer a discount to its own customer?"'
)

add_bold_paragraph('The Dependency Spiral')
doc.add_paragraph(
    "These two forces create a vicious cycle. Venues can't build direct customer relationships \u2192 "
    "they depend on aggregators for visibility \u2192 aggregators charge commissions AND insert themselves "
    "into dine-in transactions \u2192 venues lose margin on both delivery AND in-venue sales \u2192 venues "
    "can't afford to invest in their own engagement systems \u2192 dependency deepens."
)

add_bold_paragraph('The Data Asymmetry')
doc.add_paragraph(
    "Aggregators give venues dashboard access to aggregate metrics — order volumes, revenue trends, "
    "average order values, popular items, and operational scores. What they don't share is individual "
    "customer identity. A venue can see '200 orders this week' but can't identify who ordered, whether "
    "they're repeat customers, what else they eat, or when they'll come back. The platform masks customer "
    "phone numbers behind relay systems, so the venue can't contact its own customers directly. "
    "The aggregator builds a rich customer graph across the entire ecosystem; the venue gets a dashboard "
    "of its own aggregate numbers. That's not customer knowledge — that's a report card."
)

add_bold_paragraph('The Retention Gap')
doc.add_paragraph(
    "Most independent venues have zero systematic way to bring customers back. No loyalty program, "
    "no re-engagement tools, no way to reward regulars or incentivize return visits. The ones that try "
    "are stuck with punch cards or basic points systems that customers forget about. Research shows "
    "over 50% of loyalty memberships become inactive. Chains run loyalty programs, but they're expensive, "
    "generic, and disconnected from the in-venue experience."
)

add_bold_paragraph('The Result')
doc.add_paragraph(
    "Venues are caught between delivery commissions (15-28%) they can't avoid and dine-in discount "
    "programs (15-40% + commission) that cannibalise their own walk-in customers. The solution isn't to "
    "fight delivery — it provides real value. The solution is to make the in-venue experience so compelling "
    "that customers come directly, stay loyal, and keep coming back without needing aggregator-mediated "
    "discounts. That's what CheckIn builds."
)

# ── The Solution ──────────────────────────────────────────────
doc.add_heading('The Solution', level=2)

doc.add_paragraph(
    "CheckIn flips this dynamic. Instead of sending customers to an aggregator, CheckIn brings "
    "the platform to the venue — giving each venue its own intelligent engagement system."
)

doc.add_heading('For the Venue', level=3)
doc.add_paragraph(
    'CheckIn is a platform where each venue gets a deep, customizable presence — their own page, '
    'their own quests, their own brand identity — powered by CheckIn. When a diner scans a venue\'s '
    'QR code, they enter that venue\'s world: active quests, the ordering system, VCs, leaderboard '
    'position, collectibles, personalized recommendations. It feels like the venue\'s system, but it '
    'runs on CheckIn\'s platform ("Powered by CheckIn").'
)
doc.add_paragraph(
    "When the diner leaves that venue's page, they return to the broader CheckIn platform — where "
    "they can browse other venues, see what's happening around them, discover new places through GCs — "
    "but they're not 'checked in' anywhere. The check-in is the moment of commitment; the platform is "
    "the ecosystem around it."
)

p = doc.add_paragraph('The venue gets:')
add_bullet('A community of every diner who has ever checked in — auto-followed, reachable via notifications and feed, with full behavioral profiles')
add_bullet('A customer intelligence system that learns diner preferences, visit patterns, and spending behavior')
add_bullet('A gamification engine (quests, streaks, challenges) that drives repeat visits without discounting')
add_bullet('A dual credit economy that creates a genuine reason for customers to return')
add_bullet('An AI agent that can autonomously run campaigns, optimize quests, manage inventory-linked promotions, and personalize the dining experience')
add_bullet('Full ownership of customer data and relationships within the platform')

doc.add_heading('For the Diner', level=3)
doc.add_paragraph(
    "CheckIn transforms a routine meal into an engaging experience. Diners earn Venue Credits (VCs) "
    "by completing quests (try 3 new dishes, visit 3 Fridays in a row, bring 2 friends). These credits "
    "are redeemable at that venue for real value. Over time, diners also earn Global Credits (GCs) "
    "that work across all CheckIn venues, creating a discovery engine that rewards exploring new places."
)
add_bullet('A reason to go back (active quests, streaks, rewards they\'ve earned)')
add_bullet('A feed of updates from followed venues — new quests, events, specials')
add_bullet('Personalized recommendations based on their actual taste profile')
add_bullet('Social features (group quests, leaderboards, challenges with friends)')
add_bullet('A cross-venue discovery system that rewards trying new places')
add_bullet('Control over the relationship: unfollow any venue anytime, manage notifications')

doc.add_heading('The Core Engagement Loop', level=3)
doc.add_paragraph(
    "Diner visits venue \u2192 Scans QR / opens CheckIn \u2192 Checks in (auto-follows the venue) "
    "\u2192 Sees active quests, personalized recommendations \u2192 Completes quest actions (orders, "
    "visits, social) \u2192 Earns Venue Credits (VCs) \u2192 Redeems VCs for rewards at that venue "
    "\u2192 Earns Global Credits (GCs) for platform-wide activity \u2192 GCs unlock discovery of new "
    "venues \u2192 Visits new venue \u2192 Checks in \u2192 New quests begin \u2192 Cycle repeats, "
    "data compounds, experiences improve."
)
p = doc.add_paragraph()
run = p.add_run('Check-in = auto-follow. ')
run.bold = True
p.add_run(
    "Once checked in, the diner becomes part of that venue's community. The venue can reach them — "
    "notifications, feed updates, new quests, special events. If they don't want to hear from a venue "
    "anymore, they unfollow. No consent popups, no phone number sharing debates. The check-in IS the "
    "relationship. The unfollow IS the exit."
)

# ── Dual Credit Economy ──────────────────────────────────────
doc.add_heading('The Dual Credit Economy', level=2)
doc.add_paragraph(
    "CheckIn runs on a two-currency system, each serving a distinct purpose. The system draws from "
    "Free-to-Play gaming economics — the Soft Currency vs. Hard Currency model that drives retention "
    "in mobile games."
)

doc.add_heading('Venue Credits (VCs) — The "Lock-In" Currency', level=3)
add_bullet('Earned by completing quests and challenges at a specific venue')
add_bullet('Redeemable ONLY at that venue for goods, services, or experiences')
add_bullet('Non-transferable, non-exchangeable between venues')
add_bullet('Floating value: venues set their own redemption rates, enabling dynamic demand shaping')
add_bullet('Dynamic decay: 60 days of inactivity \u2192 10% balance degradation (prevents liability ballooning)')
add_bullet('Legal classification: loyalty points, exempt from VDA and PPI regulations (CBDT Notification 74/2022)')

doc.add_heading('Global Credits (GCs) — The "Liquidity" Currency', level=3)
add_bullet('Earned through platform-wide activity (cross-venue visits, referrals, brand campaigns)')
add_bullet('Redeemable at ANY CheckIn venue')
add_bullet('Fixed 1:1 INR peg (\u20b91 = 1 GC) — no floating exchange, no speculation')
add_bullet('Strict backing: brand/marketing funds deposited before GCs are minted')
add_bullet('Creates the network effect: more venues = more places to earn and redeem = more valuable for every diner')

doc.add_heading('The Swap Mechanism — "Casino Chip" Psychology', level=3)
doc.add_paragraph(
    "The GC\u2192VC conversion is the key retention mechanism. A venue offers a minting bonus "
    '("Swap GCs for our VCs today \u2192 50% bonus"). The diner swaps 1,000 GCs \u2192 receives '
    "1,500 VCs. VCs generally can't convert back to GCs (or incur a ~50% penalty). The diner is now "
    "psychologically committed to returning to that venue. This is deliberate demand-shaping — "
    "venues can incentivize swaps on slow days."
)

doc.add_heading('Three-Tier Hierarchy', level=3)
add_table(
    ['Level', 'Scope', 'Example', 'Revenue Angle'],
    [
        ['Venue VCs', 'Single independent venue', "A standalone caf\u00e9's loyalty points", 'CheckIn takes a % cut per circulation cycle'],
        ['Chain VCs', 'All outlets in a franchise/chain', '"Social VCs" redeemable at any Social outlet', 'Platform/settlement fee to franchisor'],
        ['Global Credits (GCs)', 'Every venue on CheckIn', 'Platform-wide discovery currency', 'Brand-sponsored campaigns, platform circulation'],
    ]
)
doc.add_paragraph(
    "This creates a natural upsell path: independent venues start with Venue VCs, chains upgrade to "
    "Chain VCs, and everyone participates in the GC ecosystem."
)

# ── AI-Native Architecture ───────────────────────────────────
doc.add_heading('The AI-Native Architecture', level=2)
p = doc.add_paragraph()
run = p.add_run('CheckIn is not SaaS with AI bolted on. The AI is the product.')
run.bold = True

doc.add_heading('Per-Venue AI Agent', level=3)
doc.add_paragraph("Every venue on CheckIn gets its own AI agent. This agent:")
add_bullet('', bold_prefix='Learns ')
doc.paragraphs[-1].add_run("the venue's customer base — who visits, what they order, when they come, who they come with")
add_bullet('', bold_prefix='Creates ')
doc.paragraphs[-1].add_run("quests and campaigns autonomously — tied to real business goals (move inventory, fill slow nights, reward regulars)")
add_bullet('', bold_prefix='Personalizes ')
doc.paragraphs[-1].add_run("the diner experience — recommendations based on taste profile, not popularity")
add_bullet('', bold_prefix='Optimizes ')
doc.paragraphs[-1].add_run("in real time — adjusts quest parameters based on completion rates, adjusts rewards based on margins")
add_bullet('', bold_prefix='Grows ')
doc.paragraphs[-1].add_run("the venue's engagement — without the venue owner needing to do anything after initial setup")

doc.add_heading('Technical Stack', level=3)
add_table(
    ['Component', 'Technology', 'Why'],
    [
        ['Framework', 'Vercel AI SDK 6 + Next.js', 'AI-native, built for streaming and agent orchestration'],
        ['Database', 'Supabase', 'Postgres + real-time + auth + edge functions. Zero DevOps.'],
        ['Agent Protocol', 'MCP (Model Context Protocol)', 'Each venue agent = a configuration, not a deployment'],
        ['Models', 'Multi-model routing', 'Haiku (simple) \u2192 Sonnet (recommendations) \u2192 Opus (strategy)'],
        ['Cost/venue', '\u20b9300\u20131,100/month', 'vs. \u20b93K\u201315K/month pricing = 3\u20135x margin'],
    ]
)

doc.add_heading('Agentic Revenue Model', level=3)
doc.add_paragraph(
    "Traditional SaaS charges a flat subscription. CheckIn's revenue grows as the AI does more."
)
add_table(
    ['Timeline', 'What the AI Does', 'Revenue'],
    [
        ['Month 1', 'Base platform, basic quests', '\u20b93,000/month'],
        ['Month 3', 'AI runs campaigns, personalized recommendations', '\u20b93,000 base + \u20b92,000 usage'],
        ['Month 6', 'Full autonomy, inventory-linked quests, 2 plugins', '\u20b93,000 base + \u20b94,000 usage + \u20b91,500 plugins'],
    ]
)
doc.add_paragraph(
    "Revenue correlates with value delivered. More engagement = more revenue = proof of value. "
    "This is the agentic revenue model — the AI itself generates revenue-producing actions."
)

# ── What CheckIn is NOT ──────────────────────────────────────
doc.add_heading('What CheckIn is NOT', level=2)

items = [
    ('Not an aggregator in the Zomato/Swiggy sense. ',
     'CheckIn does aggregate venues on its platform — it has to, because the Global Credits economy '
     'and cross-venue discovery only work with a network. But revenue comes from SaaS tiers and agentic value, '
     'not from taxing every order. If CheckIn provides payment gateway services, a per-transaction processing '
     'fee applies — but that\'s infrastructure cost recovery, not a commission on food sales.'),
    ('Not a delivery platform. ',
     'CheckIn is about the in-venue experience and direct customer relationships. Delivery aggregators provide '
     'genuine logistics value — CheckIn doesn\'t compete with that.'),
    ('Not a POS system. ',
     'CheckIn integrates with existing POS (Petpooja, Posist, etc.) but doesn\'t replace it.'),
    ('Not a discount platform. ',
     'This is a critical distinction from Zomato Pay / Swiggy Dineout. Those programs drive visits through '
     'mandatory 15-40% discounts. CheckIn drives visits through quests, streaks, and earned rewards.'),
    ('Not a generic loyalty tool. ',
     'Tools like Reelo do basic points-per-purchase. CheckIn\'s gamification engine drives behavior change, '
     'not just reward accumulation.'),
]
for bold_part, rest in items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bold_part)
    run.bold = True
    p.add_run(rest)

# ── Target Market ────────────────────────────────────────────
doc.add_heading('Target Market', level=2)

doc.add_heading('Primary: Urban Independent F&B Venues in India', level=3)
add_bullet('Caf\u00e9s, restaurants, bars, breweries, cloud kitchens with dine-in')
add_bullet('Located in urban clusters (starting: Indiranagar, Koramangala, HSR Layout, Bangalore)')
add_bullet('Revenue: \u20b910L\u20132Cr/month')
add_bullet('Currently have zero or basic loyalty/engagement systems')
add_bullet('Paying 15-30% to aggregators and want alternatives')
add_bullet('69% of Gen-Z consumers use aggregators for venue discovery (V1 market research)')

doc.add_heading('Secondary: F&B Chains and Franchises', level=3)
add_bullet('Multi-outlet operations wanting unified loyalty across locations')
add_bullet('Chain VCs enable cross-outlet redemption')
add_bullet('Higher ACV, longer sales cycle, but stronger retention')

doc.add_heading('Future: Any Venue with Repeat Customers', level=3)
add_bullet('Salons, gyms, co-working spaces, retail — the engagement model generalizes')
add_bullet('74% of Gen-Z consumers value real-world experiences more than digital ones (V1 research)')
add_bullet('But F&B first, because visit frequency is highest and the aggregator pain is sharpest')

# ── Product Phases ───────────────────────────────────────────
doc.add_heading('Product Phases', level=2)

doc.add_heading('Phase 1A: The Prototype — Vision Demo (Now)', level=3)
doc.add_paragraph(
    "A vibe-coded front-end prototype demonstrating the full system end-to-end. Not deployed to "
    "real venues — it's the vision made tangible for investors and venue partners."
)
p = doc.add_paragraph('Three surfaces:')
add_bullet('', bold_prefix='Diner Browser: ')
doc.paragraphs[-1].add_run('QR check-in, quest system, ordering, VC earn/redeem, gated rewards')
add_bullet('', bold_prefix='Diner App: ')
doc.paragraphs[-1].add_run('Discovery, feed, profile, GC economy, personal AI agent, collections')
add_bullet('', bold_prefix='Venue Side: ')
doc.paragraphs[-1].add_run('AI agent (WhatsApp), analytics dashboard, campaign builder, plugin showcase')

doc.add_heading('Phase 1B: The MVP — First Deployment', level=3)
doc.add_paragraph("The focused slice deployed to real venues in Indiranagar/Koramangala.")
p = doc.add_paragraph()
run = p.add_run("What's in: ")
run.bold = True
p.add_run(
    "QR-based browser check-in and auto-follow, core quest system, VC earning and redemption, "
    "basic venue dashboard, campaign builder, gated app download flow."
)
p = doc.add_paragraph()
run = p.add_run("What's NOT in (yet): ")
run.bold = True
p.add_run(
    "Full app with discovery/feed/AI agent, GC cross-venue economy, plugin ecosystem, "
    "AI-driven autonomous quest creation, brand partnerships."
)

doc.add_paragraph()
add_bold_paragraph('Key MVP Metrics')
add_table(
    ['Metric', 'What It Proves', 'Target'],
    [
        ['Quest Completion Rate', 'Are quests engaging enough?', '>40%'],
        ['Return Visit Rate', 'Are diners coming back because of CheckIn?', '>35% within 30 days'],
        ['VC Redemption Rate', 'Are credits valuable enough to use?', '>60%'],
        ['QR \u2192 Check-in Conversion', 'Does the entry flow work?', '>50% of scans'],
        ['App Download Rate', 'Does the gated reward model drive installs?', '>20% of browser users'],
        ['Venue Retention', 'Are venues staying on after trial?', '>80% after 3 months'],
        ['Venue Owner NPS', 'Do venues actually like this?', '>50'],
    ]
)

doc.add_heading('Phase 2: Intelligence & Expansion (Post-Seed, 0-12 months)', level=3)
add_bullet('Deploy app with discovery flow, feed, user profiles')
add_bullet('Per-venue AI agent learning from real customer data')
add_bullet('Automated quest creation tied to business goals')
add_bullet('Global Credits launch across venue network')
add_bullet('Expand to additional Bangalore neighborhoods')

doc.add_heading('Phase 3: Platform & Network Effects (12-24 months)', level=3)
add_bullet('Full agentic autonomy — venues set goals, AI handles execution')
add_bullet('Chain VC support, brand partnership marketplace, plugin marketplace')
add_bullet('Cross-city expansion')
add_bullet('Network effects compound — every new venue makes every existing venue more valuable')

# ── Competitive Positioning ──────────────────────────────────
doc.add_heading('Competitive Positioning', level=2)

add_table(
    ['Dimension', 'Aggregators\n(Zomato/Swiggy)', 'Dine-In Programs\n(Zomato Pay/Dineout)', 'Loyalty Tools\n(Reelo/Hashtag)', 'CheckIn'],
    [
        ['Customer\nownership', 'Aggregator owns;\nvenue gets dashboards', 'Aggregator mediates\neven walk-in transactions', 'Venue has basic\nCRM data', 'Check-in = auto-follow.\nVenue has full profile'],
        ['Cost to\nvenue', '15-28% delivery\ncommission + fees', '15-40% mandatory\ndiscount + 4-12%\ncommission', 'Flat SaaS\n(\u20b92-5K/month)', 'Tiered SaaS +\nagentic usage;\nno food commission'],
        ['Engagement\nmodel', 'Discounts and\nlisting rank', 'Discount-driven\nfootfall', 'Points-per-purchase', 'Gamified quests,\nstreaks, earned rewards'],
        ['Data depth', 'Aggregate metrics\n(volume, AOV, ratings)', 'Transaction data\nstays with platform', 'Transaction history,\nphone numbers', 'Full taste profile,\nbehavioral patterns,\nsocial graph'],
        ['Network\neffects', 'Discovery via\naggregator', 'None', 'None', 'Cross-venue\ncredit economy'],
        ['Intelligence', 'Platform-level,\nnot shared', 'None for venue', 'Basic analytics', 'Per-venue AI agent\nwith diner-level\nintelligence'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "CheckIn doesn't compete with Zomato's delivery logistics — that provides real value. CheckIn "
    "competes with the aggregator's grip on the customer relationship. By making in-venue experiences "
    "compelling enough to drive direct, repeated visits, CheckIn reduces a venue's dependency on both "
    "delivery commissions and dine-in discount programs."
)

# ── GTM ──────────────────────────────────────────────────────
doc.add_heading('Go-To-Market: Services-Led Product Discovery', level=2)
doc.add_paragraph(
    "CheckIn's GTM is deliberately manual in Phase 1. This is not a weakness — it's the strategy."
)
doc.add_paragraph(
    "The founding team does the work that AI agents will eventually automate at scale. "
    "Shash manually runs influencer campaigns. Ashish manually does venue onboarding and customer "
    "interviews while building the prototype. Aravind handles compliance and product specs."
)
doc.add_paragraph(
    "Every manual workflow is documented in structured logs. These logs become the training data and "
    "playbook for AI agents as the platform scales."
)
p = doc.add_paragraph()
run = p.add_run('The pitch: ')
run.bold = True
p.add_run(
    '"We built the entire system before asking for money. We\'ve done the work ourselves first. '
    'Now we need capital to put it in real venues and prove the economics."'
)

# ── V1 Context ───────────────────────────────────────────────
doc.add_paragraph('─' * 60)

doc.add_heading('Historical Context (V1 Thesis)', level=2)
doc.add_paragraph(
    'The original thesis (September 2025) framed CheckIn as "Building the Operating System for '
    "the Experience Economy\" — a broader framing that has been refined to focus specifically on F&B "
    "customer engagement. Key V1 concepts like the \"Identity & Trust Fabric\" (user-owned identity, "
    "bank-grade security) and the three-tiered architecture (Foundation Layer, Intelligence Layer, "
    "Application Layer) informed the current design but have evolved significantly."
)
doc.add_paragraph(
    "V1 market research provided early validation: 74% of Gen-Z consumers value real-world experiences "
    "more than digital ones; 31% of F&B executives identify customer data as a top priority but lack "
    "unified systems; over 50% of traditional loyalty memberships become inactive. These data points "
    "reinforced the focus on gamified engagement over transactional loyalty."
)

# ── Source Attribution ────────────────────────────────────────
doc.add_paragraph('─' * 60)
p = doc.add_paragraph()
p.add_run('Sources: ').italic = True
p.add_run(
    'product/overview.md (canonical, Git repo); Check-in Thesis V1 (Sep 2025); '
    'Checking V2-Working Doc; Urban Indian F&B Analysis; Investor Pitch Deck.'
).italic = True

p = doc.add_paragraph()
p.add_run(
    'This is the canonical product overview. All other documents should reference definitions '
    'and terminology from this document. If something here conflicts with another doc, this document wins.'
).italic = True

# ── Save ─────────────────────────────────────────────────────
output_path = '/Users/ashishnoel/Documents/Projects/Checkin/Checkin_Brain/final-docs/CheckIn — Product Overview.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
