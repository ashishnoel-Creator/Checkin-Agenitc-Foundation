#!/usr/bin/env python3
"""Generate CheckIn — User App PRD (Native Application).docx
Screen-by-screen flows for the native iOS/Android app experience."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Style setup ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)

def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def screen_header(name, subtitle=None):
    doc.add_heading(name, level=2)
    if subtitle:
        p = doc.add_paragraph()
        p.add_run(subtitle).italic = True

def flow_arrow(text):
    p = doc.add_paragraph()
    run = p.add_run(f'\u2192 {text}')
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x99)
    run.bold = True


# ═══════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('PRD — User Application (Native App)', level=1)

meta = doc.add_paragraph()
meta.add_run('Product: ').bold = True
meta.add_run('CheckIn Native App (iOS + Android)')
meta.add_run('\nLast updated: ').bold = True
meta.add_run('February 16, 2026')
meta.add_run('\nOwner: ').bold = True
meta.add_run('Ashish')
meta.add_run('\nPhase: ').bold = True
meta.add_run('Phase 1A (Prototype) + Phase 2 (Post-Seed deployment)')
meta.add_run('\nEntry points: ').bold = True
meta.add_run('App Store download, gated prompt from browser webapp, referral links, organic search')

doc.add_paragraph('\u2500' * 60)

# ── Overview ─────────────────────────────────────────────────
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    "The CheckIn app is where the full platform experience lives. While the browser webapp handles "
    "the in-venue check-in (scan \u2192 quest \u2192 earn), the app is the persistent relationship layer "
    "between visits. It's where diners discover new venues, manage their credits across the platform, "
    "follow their favourite spots, track streaks, and interact with their personal AI agent."
)
doc.add_paragraph(
    "The app unlocks everything the browser gates: leaderboards, collectible badges, venue feed, "
    "Global Credits economy, cross-venue discovery, and the personal AI concierge. The conversion "
    "from browser to app is the key growth metric."
)

bold_para('Core Design Principles')
bullet('The app is the relationship, the browser is the transaction')
bullet('Feed-first home screen \u2014 always something new from followed venues')
bullet('Credits are the connective tissue \u2014 VC/GC balances visible everywhere')
bullet('Discovery through rewards \u2014 GCs incentivize trying new places')
bullet('Personal AI agent as the primary navigation aid \u2014 "What should I eat tonight?"')
bullet('Social without being a social network \u2014 friend activity, group quests, leaderboards')

doc.add_paragraph('\u2500' * 60)

# ── RELATIONSHIP TO BROWSER WEBAPP ───────────────────────────
doc.add_heading('Relationship to Browser Webapp', level=2)
add_table(
    ['Aspect', 'Browser Webapp', 'Native App'],
    [
        ['When', 'During a venue visit', 'Between visits + during visits'],
        ['Entry', 'QR scan at venue', 'App icon on home screen'],
        ['Auth', 'Phone OTP (session-based)', 'Persistent login (biometric)'],
        ['Check-in', '\u2713 Full flow', '\u2713 Full flow + NFC/Bluetooth (future)'],
        ['Quests', '\u2713 Active quests at current venue', '\u2713 All quests across all venues'],
        ['Credits', '\u2713 VCs for current venue', '\u2713 Full wallet: all VCs + GCs'],
        ['Discovery', '\u2718 Gated', '\u2713 Browse venues, map, recommendations'],
        ['Feed', '\u2718 Gated', '\u2713 Updates from all followed venues'],
        ['Leaderboard', '\u2718 Gated (preview only)', '\u2713 Venue + platform rankings'],
        ['Collections', '\u2718 Gated', '\u2713 Badges, achievements, collectibles'],
        ['AI Agent', '\u2718 Not available', '\u2713 Personal concierge'],
        ['GC Economy', '\u2718 Not available', '\u2713 Earn, swap, discover via GCs'],
    ]
)

doc.add_paragraph('\u2500' * 60)

# ══════════════════════════════════════════════════════════════
# SCREEN-BY-SCREEN FLOWS
# ══════════════════════════════════════════════════════════════

doc.add_heading('Screen-by-Screen Flows', level=1)

# ── SCREEN 1: ONBOARDING ────────────────────────────────────
screen_header('Screen 1: Onboarding & First Launch',
              'First-time app open \u2014 account setup or login')

bold_para('1A: Welcome Carousel (3 screens, swipeable)')
bullet('Screen A: "Own your dining experience" \u2014 illustration of check-in + quests')
bullet('Screen B: "Earn rewards everywhere you go" \u2014 illustration of VCs + GCs flowing')
bullet('Screen C: "Your personal food concierge" \u2014 illustration of AI agent chat')
bullet('CTA: "Get Started \u2192"')

bold_para('1B: Authentication')
bullet('If user already checked in via browser: "Welcome back, {name}! We found your account." \u2192 auto-login via phone match')
bullet('If new: Phone OTP flow (identical to browser Screen 2)')
bullet('After auth: quick preference onboarding (Screen 1C)')

bold_para('1C: Taste Profile Setup (One-Time, Skippable)')
bullet('Cuisine preferences: tap to select from grid (Indian, Italian, Japanese, Chinese, Continental, Cafe, Bar, etc.)')
bullet('Dietary restrictions: Veg / Non-Veg / Vegan / Gluten-Free (multi-select)')
bullet('Spice tolerance: slider (Mild \u2192 Medium \u2192 Fire)')
bullet('Budget preference: \u20b9 / \u20b9\u20b9 / \u20b9\u20b9\u20b9 / \u20b9\u20b9\u20b9\u20b9')
bullet('"Skip for now" option \u2014 AI agent refines preferences over time from actual behavior')

bold_para('1D: Location Permission')
bullet('Prompt: "Allow CheckIn to use your location to find venues near you"')
bullet('If denied: app works but discovery defaults to city-level, not proximity')

flow_arrow('Onboarding complete \u2192 Screen 2 (Home Feed)')

# ── SCREEN 2: HOME FEED ─────────────────────────────────────
screen_header('Screen 2: Home Feed',
              'The default screen \u2014 what users see every time they open the app')

bold_para('Layout (top to bottom)')

doc.add_heading('2A: Top Bar', level=3)
bullet('CheckIn logo (top-left)')
bullet('Search icon (top-right) \u2192 Screen 6 (Search)')
bullet('Notification bell with unread count \u2192 Screen 10 (Notifications)')
bullet('Profile avatar (top-right) \u2192 Screen 5 (Profile)')

doc.add_heading('2B: Credit Summary Strip', level=3)
bullet('Horizontal card always visible below top bar:')
bullet('Left: "\U0001fa99 GC: {balance}" (tappable \u2192 Screen 4 Wallet)')
bullet('Right: "VCs at {X} venues" (tappable \u2192 Screen 4 Wallet)')
bullet('Subtle animation when credits change (coin pulse)')

doc.add_heading('2C: Active Quests Banner (if any)', level=3)
bullet('Horizontal scroll of active quest cards from followed venues')
bullet('Shows: venue logo, quest name, progress, time remaining')
bullet('"3 active quests \u2014 complete them to earn 800 VCs" (motivational)')
bullet('Tap \u2192 deep-links to quest detail for that venue')

doc.add_heading('2D: Feed (Main Content)', level=3)
doc.add_paragraph(
    "Chronological + algorithmically weighted feed of updates from followed venues:"
)
bullet('', bold_prefix='New Quest: ')
doc.paragraphs[-1].add_run('"{Venue} just launched a new quest: Taste Adventure \u2014 try 5 dishes for 300 VCs"')
bullet('', bold_prefix='Event: ')
doc.paragraphs[-1].add_run('"{Venue} is hosting Live Jazz Night this Friday \u2014 bonus 2x credits"')
bullet('', bold_prefix='Special: ')
doc.paragraphs[-1].add_run('"{Venue} has a limited-time swap bonus: GC\u2192VC at 50% extra today!"')
bullet('', bold_prefix='Streak Reminder: ')
doc.paragraphs[-1].add_run('"Your streak at {Venue} expires tomorrow \u2014 visit to keep it alive!"')
bullet('', bold_prefix='Friend Activity: ')
doc.paragraphs[-1].add_run('"Rahul just completed Loyalty Pays at Toit \u2014 he\'s now #3 on their leaderboard" (opt-in)')
bullet('', bold_prefix='New Venue Nearby: ')
doc.paragraphs[-1].add_run('"Third Wave Coffee just joined CheckIn \u2014 50 GC welcome bonus for first check-in!"')

bold_para('Feed Empty State (New User)')
bullet('"No updates yet! Follow some venues to see what\'s happening."')
bullet('CTA: "Discover venues near you \u2192" \u2192 Screen 3')
bullet('Suggested venues based on location + taste preferences')

doc.add_heading('2E: Bottom Navigation Bar', level=3)
add_table(
    ['Tab', 'Icon', 'Screen', 'Description'],
    [
        ['Home', '\U0001f3e0', 'Screen 2', 'Feed + active quests'],
        ['Discover', '\U0001f50d', 'Screen 3', 'Find new venues, map, GC opportunities'],
        ['Wallet', '\U0001fa99', 'Screen 4', 'All credits, transaction history, rewards'],
        ['Profile', '\U0001f464', 'Screen 5', 'Taste profile, collections, settings'],
        ['AI Agent', '\u2728', 'Screen 9', 'Personal concierge chat'],
    ]
)

flow_arrow('Bottom nav controls all primary navigation')

# ── SCREEN 3: DISCOVER ───────────────────────────────────────
screen_header('Screen 3: Discover',
              'Find new venues, browse by category, GC-powered discovery')

bold_para('Layout')

doc.add_heading('3A: Search Bar + Filters', level=3)
bullet('Search: "Search venues, cuisines, areas..."')
bullet('Filter chips: Cuisine type, Area, Price range, "Has active quests", "GC welcome bonus"')
bullet('Sort: Nearest / Trending / New on CheckIn / Best rewards')

doc.add_heading('3B: Map View (Toggle)', level=3)
bullet('Interactive map showing nearby CheckIn venues as pins')
bullet('Pin color indicates: \U0001f7e2 active quest available, \U0001f535 no active quest, \u2728 GC welcome bonus')
bullet('Tap pin \u2192 venue preview card (name, photo, cuisine, distance, active quests count)')
bullet('Tap card \u2192 Screen 3C (Venue Profile)')

doc.add_heading('3C: Venue List View (Default)', level=3)
doc.add_paragraph("Scrollable list of venue cards, each showing:")
bullet('Venue photo + name + cuisine tags')
bullet('Distance from current location')
bullet('Active quests count: "3 active quests"')
bullet('Top reward preview: "Earn up to 500 VCs"')
bullet('GC welcome bonus (if available): "\U0001f31f 50 GC for first check-in"')
bullet('Friend presence: "2 friends follow this venue" (if applicable)')
bullet('Tap \u2192 Screen 3D (Venue Detail Page)')

doc.add_heading('3D: Venue Detail Page (Pre-Visit)', level=3)
doc.add_paragraph(
    "Full venue page accessible from Discover. Same branded experience as the browser, "
    "but richer since they're in the app:"
)
bullet('Hero image + venue branding')
bullet('About: cuisine, hours, address, phone, Instagram link')
bullet('Active quests: full list with progress (if already following)')
bullet('Rewards available: grid of what you can earn')
bullet('Leaderboard preview: top 5 + your rank (if you\'re on it)')
bullet('Community stats: "{X} followers \u2022 {Y} check-ins this week"')
bullet('"Check In" button (if at venue, detected via GPS)')
bullet('"Follow" button (if not at venue \u2014 follow for feed updates without checking in)')
bullet('GC\u2192VC Swap Offers: if venue has an active swap bonus, prominently displayed')
bullet('Reviews/photos (future: user-generated content)')

bold_para('GC-Powered Discovery')
doc.add_paragraph(
    "The Discover tab is where Global Credits drive cross-venue exploration:"
)
bullet('"Earn 50 GC for your first check-in at any of these new venues" \u2014 discovery quests')
bullet('"Visit 3 new cafes this month \u2192 200 GC" \u2014 exploration quest')
bullet('"Swap your GCs for VCs at {Venue} \u2014 50% bonus today!" \u2014 venue-specific offers')
bullet('GC balance always visible, reminding users they have "money to spend" at new places')

flow_arrow('Tap venue \u2192 Screen 3D Venue Detail | Tap "Check In" \u2192 Browser check-in flow')

# ── SCREEN 4: WALLET ─────────────────────────────────────────
screen_header('Screen 4: Wallet & Credits',
              'Complete financial picture \u2014 all VCs, GCs, history, rewards')

bold_para('Layout')

doc.add_heading('4A: Global Credits Section', level=3)
bullet('Big balance display: "\U0001fa99 {X} Global Credits"')
bullet('Equivalent value: "\u2248 \u20b9{X}" (1:1 peg, always shown)')
bullet('"How to earn more GCs" expandable section:')
bullet('  \u2022 Visit a new venue: +50 GC', bold_prefix='  ')
bullet('  \u2022 Refer a friend: +100 GC', bold_prefix='  ')
bullet('  \u2022 Complete a Global Quest: varies', bold_prefix='  ')
bullet('"Swap GCs \u2192 VCs" button \u2192 shows venues with active swap bonuses')

doc.add_heading('4B: Venue Credits Section', level=3)
doc.add_paragraph("List of all venues where diner holds VCs:")
bullet('Each row: Venue logo + name, VC balance, expiry date')
bullet('Sorted by: balance (highest first) or recency (last visited first)')
bullet('Tap venue row \u2192 expanded view: rewards available at that venue, redemption options, transaction history')
bullet('Expiry warnings: \u26a0\ufe0f "150 VCs at Toit expire in 14 days"')
bullet('Total VC value shown at top: "VCs across {X} venues: {total} VCs"')

doc.add_heading('4C: Transaction History', level=3)
bullet('Chronological feed of all credit events:')
bullet('"Earned 200 VCs at Toit (Raise the Stakes quest)" \u2014 Feb 15')
bullet('"Redeemed 100 VCs at Toit (Free Espresso)" \u2014 Feb 14')
bullet('"Earned 50 GC (First check-in at Third Wave)" \u2014 Feb 12')
bullet('"Swapped 500 GC \u2192 750 VCs at Arbor Brewing (50% bonus)" \u2014 Feb 10')
bullet('Filter by: venue, credit type (VC/GC), earned/spent, date range')

doc.add_heading('4D: GC \u2192 VC Swap Flow', level=3)
doc.add_paragraph("When diner taps 'Swap GCs' or a venue's swap offer:")
bullet('Step 1: Select venue to swap into')
bullet('Step 2: Enter GC amount to swap')
bullet('Step 3: Preview: "500 GC \u2192 750 VCs at {Venue} (50% bonus!)"')
bullet('Step 4: Confirm swap')
bullet('Step 5: Celebration animation. VCs appear in venue wallet instantly.')
bullet('Warning: "VCs at {Venue} can\'t be converted back to GCs" (clear, no surprises)')

bold_para('Edge Cases')
bullet('No GCs: "Earn Global Credits by visiting new venues or referring friends" with CTA to Discover')
bullet('VC decay warning: "Your VCs at {Venue} will start decreasing in 15 days due to inactivity"')
bullet('Venue no longer on CheckIn: "This venue is no longer active. Your remaining VCs have been converted to GCs at 50% value." (safety net)')

flow_arrow('Tap venue \u2192 Venue rewards | Tap swap \u2192 Swap flow')

# ── SCREEN 5: PROFILE ────────────────────────────────────────
screen_header('Screen 5: Profile & Collections',
              'Diner identity, achievements, taste profile, settings')

bold_para('Layout')

doc.add_heading('5A: Profile Header', level=3)
bullet('Profile photo, name, member since date')
bullet('Headline stats: "{X} check-ins \u2022 {Y} venues \u2022 {Z} quests completed"')
bullet('"Edit Profile" button')

doc.add_heading('5B: Taste Profile Card', level=3)
bullet('Visual taste map: radar chart showing cuisine preferences (auto-generated from actual order data + initial setup)')
bullet('Top cuisines: "Indian \u2022 Italian \u2022 Japanese"')
bullet('Dietary: "Non-Veg \u2022 No restrictions"')
bullet('Spice: "Medium-High \U0001f336\ufe0f"')
bullet('Budget: "\u20b9\u20b9\u20b9"')
bullet('AI-refined: "Based on your last 23 check-ins and 47 orders"')
bullet('Editable: diner can override any AI-generated preference')

doc.add_heading('5C: Collections & Badges', level=3)
doc.add_paragraph("Grid of earned badges and collectibles:")
bullet('', bold_prefix='Venue Badges: ')
doc.paragraphs[-1].add_run('"Toit Regular" (10+ visits), "Third Wave Explorer" (first visit), "Arbor VIP" (500+ VCs earned)')
bullet('', bold_prefix='Quest Badges: ')
doc.paragraphs[-1].add_run('"Taste Adventurer" (try 20 unique dishes), "Social Butterfly" (refer 5 friends), "Streak Master" (7-day streak)')
bullet('', bold_prefix='Platform Badges: ')
doc.paragraphs[-1].add_run('"Early Adopter", "10 Venues", "100 Quests", "GC Whale" (1000+ GC earned)')
bullet('Locked badges visible as silhouettes with "How to earn" explanation')
bullet('Shareable: tap badge \u2192 share to Instagram Stories / WhatsApp')

doc.add_heading('5D: Leaderboard Summary', level=3)
bullet('Per-venue rankings: "Toit: #12 \u2022 Third Wave: #3 \u2022 Arbor: #45"')
bullet('Platform ranking: "#234 overall"')
bullet('Tap \u2192 full leaderboard (Screen 8)')

doc.add_heading('5E: Visit History', level=3)
bullet('Timeline of all check-ins with dates, venues, credits earned')
bullet('Streak tracker: "Current streak: 3 weeks at Toit \U0001f525"')
bullet('Monthly summary: "Feb 2026: 8 check-ins, 4 venues, 1,200 VCs earned"')

doc.add_heading('5F: Settings', level=3)
bullet('Notification preferences (per venue: on/off, frequency)')
bullet('Privacy: friend activity visibility (show/hide)')
bullet('Data export: download full profile data (GDPR-aligned)')
bullet('Delete account: clear all data, forfeit credits')
bullet('Linked accounts: phone, Google')
bullet('Help & support, Terms, Privacy Policy')

flow_arrow('Edit profile \u2192 Edit flow | Tap badge \u2192 Share | Tap leaderboard \u2192 Screen 8')

# ── SCREEN 6: SEARCH ─────────────────────────────────────────
screen_header('Screen 6: Search',
              'Universal search across venues, quests, friends')

bold_para('Search Functionality')
bullet('Single search bar: "Search venues, quests, friends, cuisines..."')
bullet('', bold_prefix='Venue results: ')
doc.paragraphs[-1].add_run('Name, cuisine, area matches')
bullet('', bold_prefix='Quest results: ')
doc.paragraphs[-1].add_run('"Taste Adventure" shows all venues with that quest active')
bullet('', bold_prefix='Area results: ')
doc.paragraphs[-1].add_run('"Koramangala" shows all venues in that area')
bullet('', bold_prefix='Friend results: ')
doc.paragraphs[-1].add_run('Find friends on CheckIn by name')

bold_para('Recent Searches + Trending')
bullet('Recent: last 5 searches')
bullet('Trending: "Trending in Bangalore: Toit, Third Wave, Arbor Brewing"')
bullet('AI suggestion (Phase 2): "Based on your taste: try searching for \'Japanese Koramangala\'"')

flow_arrow('Tap result \u2192 Venue Detail / Quest Detail / Friend Profile')

# ── SCREEN 7: VENUE PAGE (IN-APP) ───────────────────────────
screen_header('Screen 7: Venue Page (In-App, During Visit)',
              'When user checks in via the app instead of browser')

doc.add_paragraph(
    "When a diner checks in via the app (either scanning QR or tapping 'Check In' on a venue page "
    "while at the venue), the in-app experience mirrors the browser webapp but with full app features unlocked:"
)
bullet('Same quest dashboard as browser Screen 4, but with full leaderboard access')
bullet('Real-time table pot with push notifications when group members earn credits')
bullet('Personal AI recommendations: "Based on your taste profile, try the Mushroom Risotto"')
bullet('Bill splitting with in-app payment tracking')
bullet('Badge/collectible unlocks animate directly in-app')
bullet('Post-visit: no app download prompt needed (they\'re already in the app)')
bullet('Post-visit: streak tracker, next-quest suggestion, GC earning opportunities')

bold_para('Key Difference from Browser')
doc.add_paragraph(
    "The app check-in is seamless \u2014 no OTP, no sign-up friction. The diner is already "
    "authenticated. GPS detects venue proximity, one tap to check in. Everything syncs "
    "in real-time with the browser sessions of other table members."
)

# ── SCREEN 8: LEADERBOARD ───────────────────────────────────
screen_header('Screen 8: Leaderboard',
              'Venue-level and platform-level competitive rankings')

doc.add_heading('8A: Venue Leaderboard', level=3)
bullet('Rankings within a single venue based on Stars earned')
bullet('Top 10 displayed with avatar, name, Stars count')
bullet('Your position highlighted (even if outside top 10)')
bullet('Time filter: This Week / This Month / All Time')
bullet('"How Stars work: Stars are earned alongside VCs but can\'t be spent. They\'re your status."')

doc.add_heading('8B: Platform Leaderboard', level=3)
bullet('Rankings across all CheckIn venues')
bullet('Category filters: Overall / Most Venues Visited / Longest Streak / Most Quests')
bullet('City filter: "Bangalore" (future: cross-city)')
bullet('Friend filter: "Among your friends"')

doc.add_heading('8C: Gamification Elements', level=3)
bullet('Tier system: Bronze \u2192 Silver \u2192 Gold \u2192 Platinum (based on cumulative Stars)')
bullet('Tier benefits: increased VC earning rates, exclusive quests, early access to new venues')
bullet('Weekly reset for "This Week" board creates recurring engagement')
bullet('Animated rank-up celebrations when user moves up tiers')

# ── SCREEN 9: AI AGENT ──────────────────────────────────────
screen_header('Screen 9: Personal AI Agent',
              'The AI concierge \u2014 the most differentiated feature')

bold_para('Interface')
bullet('Chat-style interface (like iMessage/WhatsApp)')
bullet('Always accessible via bottom nav "\u2728" tab')
bullet('Agent has a personality: friendly, knowledgeable about food, slightly witty')
bullet('Agent avatar: CheckIn branded character')

bold_para('What the Agent Can Do')

doc.add_heading('Venue Recommendations', level=3)
bullet('User: "Where should I eat tonight?"')
bullet('Agent: "Based on your love of Italian and your \u20b9\u20b9 budget, I\'d suggest Toscano in Indiranagar. They have a Taste Adventure quest running \u2014 try 3 pastas, earn 300 VCs. Plus you have 200 GCs you could swap for 300 VCs there with today\'s bonus!"')

doc.add_heading('Quest Guidance', level=3)
bullet('User: "What quests can I complete today?"')
bullet('Agent: "You have 3 active quests: Loyalty Pays at Toit (1 more visit needed \u2014 your streak expires tomorrow!), Taste Adventure at Arbor (2 more dishes), and a Global Quest (visit any new venue for 50 GC)."')

doc.add_heading('Credit Optimization', level=3)
bullet('User: "I have 500 GCs, what should I do with them?"')
bullet('Agent: "Arbor Brewing has a 50% swap bonus today \u2014 your 500 GCs would become 750 VCs. That\'s enough for a free craft beer flight! Otherwise, Third Wave has a 30% bonus. Or keep them liquid if you\'re exploring this weekend."')

doc.add_heading('Group Planning', level=3)
bullet('User: "Planning dinner for 6 people, mix of veg and non-veg"')
bullet('Agent: "For a group of 6 with mixed preferences, I\'d suggest Truffles (great for both, Group Quest running: table spends \u20b96K \u2192 everyone gets 300 VCs), or Sahib Sindh Sultan (amazing veg options alongside non-veg). Want me to check if any of your friends follow these?"')

doc.add_heading('Morning Brief (Push Notification)', level=3)
bullet('Daily personalized update (opt-in): "Good morning, Ashish! Your Toit streak hits 4 weeks tomorrow. Third Wave has a new quest today. And you have 150 VCs expiring in 5 days at Arbor."')

bold_para('Limitations (MVP / Phase 2)')
bullet('Agent cannot make reservations (future integration)')
bullet('Agent cannot place orders')
bullet('Agent is text-only (no voice in MVP)')
bullet('Agent recommendations limited to CheckIn venues')

bold_para('Data Sources for Agent')
bullet('Diner\'s taste profile + visit history + credit balances')
bullet('Venue quest data + reward availability + swap bonuses')
bullet('Friend activity (opt-in)')
bullet('Location + time of day')
bullet('Previous agent conversations (context memory)')

# ── SCREEN 10: NOTIFICATIONS ────────────────────────────────
screen_header('Screen 10: Notifications',
              'All alerts and updates in one place')

bold_para('Notification Types')
add_table(
    ['Type', 'Example', 'Priority'],
    [
        ['Quest progress', '"1 more visit to complete Loyalty Pays at Toit!"', 'High'],
        ['Streak warning', '"Your streak expires tomorrow \u2014 visit Toit!"', 'High'],
        ['Credit expiry', '"150 VCs at Arbor expiring in 5 days"', 'High'],
        ['New quest', '"{Venue} launched a new quest: Beat the Clock"', 'Medium'],
        ['Swap bonus', '"{Venue} is offering 50% GC\u2192VC bonus today"', 'Medium'],
        ['Friend activity', '"Rahul just hit #1 on Toit\'s leaderboard"', 'Low'],
        ['New venue', '"Third Wave Coffee just joined CheckIn near you"', 'Low'],
        ['Event', '"{Venue} is hosting Live Jazz Night this Friday"', 'Low'],
        ['AI Morning Brief', '"Your daily CheckIn summary"', 'Medium'],
        ['Badge earned', '"You just earned the Streak Master badge!"', 'Medium'],
    ]
)

bold_para('Notification Preferences')
bullet('Per-venue: on/off toggle for each followed venue')
bullet('Per-type: on/off for each notification category')
bullet('Quiet hours: no notifications between 11pm\u20137am (configurable)')
bullet('AI Morning Brief: on/off, configurable time')

doc.add_paragraph('\u2500' * 60)

# ── SCREEN MAP ───────────────────────────────────────────────
doc.add_heading('Screen Map (Quick Reference)', level=2)
add_table(
    ['Screen', 'Name', 'Purpose', 'Bottom Nav Tab'],
    [
        ['1', 'Onboarding', 'First launch, auth, taste setup', '\u2014'],
        ['2', 'Home Feed', 'Default screen, venue updates, active quests', '\U0001f3e0 Home'],
        ['3', 'Discover', 'Find venues, map, GC discovery', '\U0001f50d Discover'],
        ['3D', 'Venue Detail', 'Pre-visit venue info + follow/check-in', '\U0001f50d Discover'],
        ['4', 'Wallet', 'GCs + all VCs + history + swap', '\U0001fa99 Wallet'],
        ['5', 'Profile', 'Taste profile, collections, badges, settings', '\U0001f464 Profile'],
        ['6', 'Search', 'Universal search', 'Via search icon'],
        ['7', 'Venue (In-Visit)', 'Full quest experience during visit', 'Via check-in'],
        ['8', 'Leaderboard', 'Venue + platform rankings', 'Via profile/venue'],
        ['9', 'AI Agent', 'Personal concierge chat', '\u2728 AI'],
        ['10', 'Notifications', 'All alerts', 'Via bell icon'],
    ]
)

doc.add_paragraph('\u2500' * 60)

# ── SUCCESS METRICS ──────────────────────────────────────────
doc.add_heading('Success Metrics', level=2)
add_table(
    ['Metric', 'Target', 'Measurement'],
    [
        ['Browser \u2192 App conversion', '>20%', 'Browser users who download app within 30 days'],
        ['DAU / MAU', '>30%', 'Daily active / monthly active ratio'],
        ['Feed engagement', '>50%', 'Users who interact with at least 1 feed item per session'],
        ['Discovery conversion', '>15%', 'Discover tab users who check in at a new venue within 7 days'],
        ['GC \u2192 VC swap rate', '>40%', 'GCs earned that get swapped to VCs within 60 days'],
        ['AI Agent usage', '>25%', 'Weekly active users who chat with agent at least once'],
        ['Notification opt-in', '>60%', 'Users who keep push notifications enabled'],
        ['Collections engagement', '>30%', 'Users who view their badges at least monthly'],
        ['Multi-venue users', '>25%', 'Users who check in at 2+ venues within 30 days'],
        ['Retention D30', '>40%', 'Users who open app at least once 30 days after install'],
    ]
)

doc.add_paragraph('\u2500' * 60)

# ── Sources ──────────────────────────────────────────────────
p = doc.add_paragraph()
p.add_run('Sources: ').italic = True
p.add_run(
    'product/features.md (Phase 1A app features, V1.5, V2.0), product/overview.md (core loop, '
    'agentic vision), product/roadmap.md (Phase 2 deployment), product/credit-system.md (GC economy, '
    'swap mechanism), product/architecture.md (tech stack).'
).italic = True

# ── Save ─────────────────────────────────────────────────────
output_path = '/Users/ashishnoel/Documents/Projects/Checkin/Checkin_Brain/final-docs/CheckIn — User App PRD.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
