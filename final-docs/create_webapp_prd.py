#!/usr/bin/env python3
"""Generate CheckIn — Webapp PRD (Browser Check-In Experience).docx
Screen-by-screen flows for the venue check-in browser experience (Phase 1B MVP)."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Style setup ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = 'Calibri'
    h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    if level == 1:
        h.font.size = Pt(22)
        h.paragraph_format.space_before = Pt(24)
    elif level == 2:
        h.font.size = Pt(16)
        h.paragraph_format.space_before = Pt(18)
    else:
        h.font.size = Pt(13)
        h.paragraph_format.space_before = Pt(12)


def bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
    return table

def screen_header(name, subtitle=None):
    """Add a screen section header with optional subtitle."""
    doc.add_heading(name, level=2)
    if subtitle:
        p = doc.add_paragraph()
        p.add_run(subtitle).italic = True

def state_box(title, items):
    """Add a state/component description."""
    p = doc.add_paragraph()
    run = p.add_run(f'{title}:')
    run.bold = True
    for item in items:
        bullet(item)

def flow_arrow(text):
    """Add a flow transition line."""
    p = doc.add_paragraph()
    run = p.add_run(f'\u2192 {text}')
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x99)
    run.bold = True


# ═══════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ═══════════════════════════════════════════════════════════════

doc.add_heading('PRD — Webapp (Venue Check-In Experience)', level=1)

meta = doc.add_paragraph()
meta.add_run('Product: ').bold = True
meta.add_run('CheckIn Browser Webapp (Phase 1B MVP)')
meta.add_run('\nLast updated: ').bold = True
meta.add_run('February 16, 2026')
meta.add_run('\nOwner: ').bold = True
meta.add_run('Ashish')
meta.add_run('\nPlatform: ').bold = True
meta.add_run('Mobile browser (PWA-ready, no app install required)')
meta.add_run('\nEntry point: ').bold = True
meta.add_run('QR code at venue (table tent, receipt, bar counter, entrance)')

doc.add_paragraph('\u2500' * 60)

# ── Overview ─────────────────────────────────────────────────
doc.add_heading('Overview', level=2)
doc.add_paragraph(
    "The Webapp is CheckIn's primary entry point — the first thing a diner interacts with. "
    "It runs entirely in the mobile browser (no app install required). The diner scans a QR code "
    "at a venue, checks in, sees quests, earns credits, redeems rewards, and splits bills — all "
    "without downloading anything."
)
doc.add_paragraph(
    "The browser experience is deliberately designed to be the hook. Premium features (leaderboards, "
    "collectibles, personal AI agent, discovery feed) are visible but gated behind app download. "
    "This creates a natural conversion funnel: scan \u2192 check-in \u2192 engage \u2192 download app."
)

bold_para('Core Design Principles')
bullet('Zero friction to first value — check-in in under 15 seconds')
bullet('Mobile-first, thumb-friendly — designed for one-handed use at a table')
bullet('Venue-branded — feels like the venue\'s system, "Powered by CheckIn"')
bullet('Gated progression — browser is free, app unlocks the full experience')
bullet('Group-native — bill splitting, group quests, table pot are core, not afterthoughts')

doc.add_paragraph('\u2500' * 60)

# ══════════════════════════════════════════════════════════════
# SCREEN-BY-SCREEN FLOWS
# ══════════════════════════════════════════════════════════════

doc.add_heading('Screen-by-Screen Flows', level=1)

# ──────────────────────────────────────────────────────────────
# SCREEN 1: QR SCAN → LANDING
# ──────────────────────────────────────────────────────────────
screen_header('Screen 1: QR Scan \u2192 Venue Landing Page',
              'Entry point — first impression of CheckIn')

bold_para('Trigger')
doc.add_paragraph(
    "Diner scans a QR code placed at the venue (table tent, bar counter, entrance, receipt). "
    "QR encodes a URL: checkin.app/{venue-slug}?table={table-id}"
)

bold_para('What the diner sees')
bullet('Venue hero image / banner (full-width, venue-branded)')
bullet('Venue name, logo, and tagline')
bullet('"Powered by CheckIn" badge (subtle, bottom corner)')
bullet('Live pulse indicator: "{X} people checked in right now" (social proof)')
bullet('Primary CTA button: "Check In \u2192" (large, prominent, venue accent color)')
bullet('Secondary info: venue address, hours, cuisine tags')
bullet('Preview peek: "3 active quests \u2022 Earn up to 500 credits today" (teaser below CTA)')

bold_para('States')
bullet('', bold_prefix='New visitor: ')
doc.paragraphs[-1].add_run('Shows "Check In \u2192" button. No account yet.')
bullet('', bold_prefix='Returning visitor (not checked in): ')
doc.paragraphs[-1].add_run('Shows "Welcome back, {name}! Check In \u2192" with their avatar. Auto-detected via browser cookie/local storage.')
bullet('', bold_prefix='Already checked in: ')
doc.paragraphs[-1].add_run('Redirects straight to Screen 4 (Quest Dashboard). Shows "You\'re checked in \u2713"')

bold_para('Edge Cases')
bullet('QR scanned outside venue (GPS mismatch): Show venue page but disable check-in. Message: "Visit {venue} to check in and start earning!"')
bullet('Venue is closed: Show "Opens at {time}" with option to follow venue for updates')
bullet('Invalid/expired QR: Redirect to CheckIn homepage with venue search')

flow_arrow('Diner taps "Check In" \u2192 Screen 2')

# ──────────────────────────────────────────────────────────────
# SCREEN 2: AUTHENTICATION
# ──────────────────────────────────────────────────────────────
screen_header('Screen 2: Authentication (Sign Up / Sign In)',
              'Minimal friction — phone OTP only for MVP')

bold_para('What the diner sees')
bullet('Phone number input field (auto-detected country code +91)')
bullet('"Continue" button')
bullet('Fine print: "By continuing, you follow {venue name} and agree to CheckIn Terms"')
bullet('Alternative: "Sign in with Google" button (secondary option)')

bold_para('Flow: Phone OTP')
bullet('Step 1: Diner enters phone number \u2192 taps "Continue"')
bullet('Step 2: OTP input screen (4-digit code, auto-read from SMS)')
bullet('Step 3: If new user \u2192 Screen 2B (Name Entry). If returning user \u2192 Screen 3 (Check-In Confirmation)')

bold_para('Screen 2B: Name Entry (New Users Only)')
bullet('First name input (required)')
bullet('Optional: profile photo upload (can skip)')
bullet('"Start Exploring \u2192" button')
bullet('This screen appears only once, ever. Returning users skip it.')

bold_para('What happens in the background')
bullet('Supabase Auth creates/verifies user session')
bullet('GPS validation confirms diner is physically at the venue')
bullet('Auto-follow: diner profile linked to venue in Community Graph')
bullet('Welcome Credits: new users get a small VC bonus (configurable per venue, e.g., 50 VCs)')

bold_para('Edge Cases')
bullet('OTP not received: "Resend" button appears after 30 seconds. Fallback: "Try Google Sign-In instead"')
bullet('GPS permission denied: Allow check-in but flag it. Venue staff can manually verify. Message: "For the best experience, enable location"')
bullet('Phone number already linked to account: Auto-detect, skip name entry, welcome back')

flow_arrow('Auth complete \u2192 Screen 3')

# ──────────────────────────────────────────────────────────────
# SCREEN 3: CHECK-IN CONFIRMATION
# ──────────────────────────────────────────────────────────────
screen_header('Screen 3: Check-In Confirmation',
              'The "moment of commitment" — auto-follow happens here')

bold_para('What the diner sees (celebration micro-interaction, 2-3 seconds)')
bullet('Animated check-in confirmation: "\u2713 You\'re checked in at {Venue Name}!"')
bullet('Confetti / particle animation (brief, delightful)')
bullet('Welcome credits awarded: "+50 VCs \u2014 Welcome Aboard!" (with coin animation)')
bullet('Auto-follow confirmation: "You now follow {Venue}. You\'ll see their updates in your feed."')
bullet('Quick stats: "You\'re visitor #{X} today" (social proof)')

bold_para('Auto-transitions after 2-3 seconds to Screen 4')
doc.add_paragraph(
    "This screen is intentionally brief — it's a celebration moment, not a decision point. "
    "The diner shouldn't have to tap anything. The animation plays, credits are awarded, and "
    "the screen transitions to the Quest Dashboard."
)

bold_para('What happens in the background')
bullet('Check-in event recorded in database (timestamp, venue, table, GPS)')
bullet('Auto-follow: venue appears in diner\'s followed venues list')
bullet('Welcome Quest triggered (if first visit)')
bullet('Venue dashboard updates in real-time: new check-in notification')
bullet('If table ID was in QR: diner is associated with that table for bill/group features')

flow_arrow('Auto-transition \u2192 Screen 4')

# ──────────────────────────────────────────────────────────────
# SCREEN 4: QUEST DASHBOARD
# ──────────────────────────────────────────────────────────────
screen_header('Screen 4: Quest Dashboard (Main Hub)',
              'The primary screen during a visit — quests, credits, actions')

bold_para('Layout (top to bottom)')

doc.add_heading('4A: Header Bar', level=3)
bullet('Venue name + logo (top-left)')
bullet('Diner avatar + name (top-right, tappable \u2192 profile)')
bullet('Credit balance pill: "\U0001fa99 {X} VCs" (always visible, tappable \u2192 Screen 7)')

doc.add_heading('4B: Active Quests Section', level=3)
doc.add_paragraph(
    "Horizontally scrollable quest cards. Each card shows:"
)
bullet('Quest name and icon (e.g., "\U0001f525 Loyalty Pays", "\U0001f355 Taste Adventure")')
bullet('Progress bar: "2/3 visits completed"')
bullet('Reward preview: "Earn 200 VCs"')
bullet('Time remaining (if time-limited): "Ends in 2h 15m"')
bullet('Tap on card \u2192 Screen 5 (Quest Detail)')

bold_para('Quest Types in MVP')
add_table(
    ['Quest Type', 'Name', 'Mechanic', 'Example'],
    [
        ['First Visit', 'Welcome Aboard', 'Auto-complete on check-in', 'Check in \u2192 50 VCs'],
        ['Bill Value', 'Raise the Stakes', 'Min spend threshold', 'Spend \u20b92,000 \u2192 200 VCs'],
        ['Visit Frequency', 'Loyalty Pays', 'X visits in Y days', '3 Fridays in a row \u2192 free cocktail'],
        ['Time-Based', 'Beat the Clock', 'Check in during window', 'Happy hour check-in \u2192 100 VCs'],
        ['Menu Exploration', 'Taste Adventure', 'Try X different items', 'Try 5 dishes \u2192 secret menu item'],
        ['Social / Referral', 'Squad Mode', 'Bring X friends who check in', 'Bring 2 friends \u2192 300 VCs each'],
        ['Group Quest', 'Table Challenge', 'Group completes together', 'Table spends \u20b95K \u2192 pool gets 500 VCs'],
    ]
)

doc.add_heading('4C: Table Pot (Group Feature)', level=3)
doc.add_paragraph(
    "If multiple people are checked in at the same table:"
)
bullet('Shared "Table Pot" display: "\U0001fa99 Table Pot: 350 VCs (3 people)"')
bullet('Credits earned during this session pool into the pot')
bullet('Pot splits equally at bill time (or can be used toward bill)')
bullet('Animated coin-drop when someone earns credits ("Rahul earned +100 VCs for the table!")')

doc.add_heading('4D: Quick Actions Bar (Bottom)', level=3)
bullet('', bold_prefix='\U0001f4cb Menu: ')
doc.paragraphs[-1].add_run('View menu with recommendations \u2192 Screen 6')
bullet('', bold_prefix='\U0001fa99 Rewards: ')
doc.paragraphs[-1].add_run('See and redeem rewards \u2192 Screen 7')
bullet('', bold_prefix='\U0001f4b0 Bill: ')
doc.paragraphs[-1].add_run('View/split bill \u2192 Screen 8')
bullet('', bold_prefix='\U0001f3c6 Leaderboard: ')
doc.paragraphs[-1].add_run('Venue leaderboard (visible but gated \u2192 prompts app download)')

bold_para('Gated Features (Visible but Locked)')
doc.add_paragraph(
    "These are visible on the quest dashboard as greyed-out / locked cards with an app download prompt:"
)
bullet('\U0001f512 Leaderboard position: "Download the app to see your rank"')
bullet('\U0001f512 Collections & Badges: "Unlock collectibles in the app"')
bullet('\U0001f512 Venue Feed: "Get updates from {venue} in the app"')
bullet('\U0001f512 Discovery: "Find new venues with Global Credits in the app"')

flow_arrow('Tap quest card \u2192 Screen 5 | Tap Menu \u2192 Screen 6 | Tap Rewards \u2192 Screen 7 | Tap Bill \u2192 Screen 8')

# ──────────────────────────────────────────────────────────────
# SCREEN 5: QUEST DETAIL
# ──────────────────────────────────────────────────────────────
screen_header('Screen 5: Quest Detail',
              'Expanded view of a single quest with progress and rules')

bold_para('What the diner sees')
bullet('Quest name, icon, and description')
bullet('Full progress tracker: visual steps (e.g., 3 circles, 2 filled)')
bullet('Reward details: "Complete this quest to earn 200 VCs + \u2b50 50 Stars"')
bullet('Quest rules / fine print: "Visit must be at least 30 minutes. Bill must be \u20b9500+."')
bullet('Time remaining (if applicable)')
bullet('Quest history: "You completed this quest 2 times before"')

bold_para('Validation Methods (per quest type)')
add_table(
    ['Quest Type', 'How It\'s Validated', 'Fallback'],
    [
        ['Welcome Aboard', 'Auto-complete on check-in', 'N/A'],
        ['Raise the Stakes', 'POS bill data (if integrated) OR staff confirms bill total', 'Diner enters bill total, staff approves via dashboard'],
        ['Loyalty Pays', 'System tracks check-in history automatically', 'N/A'],
        ['Beat the Clock', 'Check-in timestamp within time window', 'N/A'],
        ['Taste Adventure', 'POS order data (if integrated) OR diner self-reports + staff confirms', 'Photo upload of dishes (future)'],
        ['Squad Mode', 'Friends\' check-in detected at same venue/table', 'Referral code entered by friend at check-in'],
        ['Group Quest', 'Combined table bill total from POS or manual entry', 'Staff confirms via dashboard'],
    ]
)

bold_para('Quest Completion State')
bullet('Animation: quest card turns gold, confetti, credit coins fly to balance')
bullet('"+{X} VCs earned!" notification')
bullet('If Stars earned: "+{X} Stars \u2014 you\'re now #{rank} on the leaderboard!" (with app download prompt to see full board)')
bullet('Next quest suggestion: "Up next: Taste Adventure \u2014 try 5 new dishes for 300 VCs"')

flow_arrow('Back \u2192 Screen 4')

# ──────────────────────────────────────────────────────────────
# SCREEN 6: MENU & ORDERING
# ──────────────────────────────────────────────────────────────
screen_header('Screen 6: Menu & Recommendations',
              'Browse menu, see AI recommendations, track quest-relevant items')

bold_para('What the diner sees')
bullet('Full venue menu organized by category (Starters, Mains, Drinks, Desserts)')
bullet('Each item: name, price, photo (if available), dietary tags (V, VG, GF)')
bullet('', bold_prefix='\u2728 Quest Highlight: ')
doc.paragraphs[-1].add_run('Items that count toward active quests are badged. E.g., if "Taste Adventure" quest is active, untried dishes show a "\U0001f31f Try this for your quest!" badge.')
bullet('', bold_prefix='\U0001f916 AI Recommendation (Phase 2): ')
doc.paragraphs[-1].add_run('"Based on your taste profile: try the Truffle Mushroom Risotto" \u2014 greyed out in MVP, shows "Unlock personalized picks in the app"')

bold_para('Ordering (MVP Scope)')
doc.add_paragraph(
    "In MVP, the menu is browse-only — ordering happens through the venue's existing process "
    "(waiter, POS, separate QR ordering system). CheckIn's menu is for discovery and quest tracking."
)
doc.add_paragraph(
    "Future (Phase 2+): In-app ordering with cart, table-side submission to POS/KDS. "
    "This requires deep POS integration and is deferred."
)

bold_para('Edge Cases')
bullet('Menu not uploaded by venue: Show "Menu coming soon. Ask your server!" with venue contact')
bullet('Item out of stock (if POS-integrated): Grey out with "Sold out" badge')

flow_arrow('Back \u2192 Screen 4')

# ──────────────────────────────────────────────────────────────
# SCREEN 7: REWARDS & WALLET
# ──────────────────────────────────────────────────────────────
screen_header('Screen 7: Rewards & Wallet',
              'VC balance, available rewards, redemption flow')

bold_para('Layout')

doc.add_heading('7A: Wallet Summary', level=3)
bullet('VC balance: "\U0001fa99 1,250 Venue Credits at {Venue Name}"')
bullet('Stars balance: "\u2b50 340 Stars \u2014 Rank #12" (with "See leaderboard in app" link)')
bullet('Expiry warning (if applicable): "150 VCs expire in 14 days \u2014 use them!"')
bullet('Transaction history link: "View history" \u2192 list of earned/redeemed credits')

doc.add_heading('7B: Available Rewards', level=3)
doc.add_paragraph("Grid of reward cards, each showing:")
bullet('Reward name: "Free Espresso", "10% Off Next Visit", "Secret Menu Access"')
bullet('Cost in VCs: "\U0001fa99 500 VCs"')
bullet('Availability: "Available now" or "Need 200 more VCs"')
bullet('Tappable: if affordable \u2192 Redemption Confirmation (Screen 7C)')

doc.add_heading('7C: Redemption Confirmation', level=3)
bullet('Modal/bottom sheet: "Redeem {Reward} for {X} VCs?"')
bullet('New balance preview: "Balance after: {remaining} VCs"')
bullet('"Confirm Redemption" button')
bullet('On confirm: QR code or alphanumeric code generated for staff to verify')
bullet('Code valid for: 15 minutes (configurable per venue)')
bullet('Staff scans/enters code on venue dashboard \u2192 reward fulfilled')

bold_para('Edge Cases')
bullet('Insufficient VCs: Card shows "Need {X} more VCs" with suggestion to complete a quest')
bullet('Reward sold out (limited quantity): "Sold out \u2014 check back tomorrow"')
bullet('Redemption code expired: "Code expired. Tap to generate a new one" (VCs refunded)')

flow_arrow('Back \u2192 Screen 4')

# ──────────────────────────────────────────────────────────────
# SCREEN 8: BILL & PAYMENT
# ──────────────────────────────────────────────────────────────
screen_header('Screen 8: Bill & Payment',
              'Bill entry, table pot usage, credit redemption, bill splitting')

bold_para('Flow')

doc.add_heading('8A: Bill Entry', level=3)
doc.add_paragraph(
    "The bill amount is entered into CheckIn so credits can be applied and quests validated."
)
bullet('Option 1 (POS-integrated): Bill auto-populated from POS system. Diner sees itemized bill.')
bullet('Option 2 (Manual, MVP default): Payer enters total bill amount manually. "Your bill: \u20b9____"')
bullet('Staff verification: venue dashboard shows pending bill for staff to confirm amount')

doc.add_heading('8B: Apply Credits', level=3)
bullet('Table Pot display: "Table Pot: 350 VCs available"')
bullet('Slider or input: "Use VCs toward this bill: {amount}"')
bullet('Real-time calculation: "Bill: \u20b93,000 \u2212 \u20b9200 (VCs) = \u20b92,800 to pay"')
bullet('VC coverage cap: venue sets max % of bill payable via VCs (e.g., 30%)')
bullet('"Apply" button confirms credit usage')

doc.add_heading('8C: Bill Split', level=3)
doc.add_paragraph(
    "If multiple people are checked in at the table:"
)
bullet('', bold_prefix='Equal split: ')
doc.paragraphs[-1].add_run('"Split equally between {X} people" \u2192 shows per-person amount')
bullet('', bold_prefix='Custom split: ')
doc.paragraphs[-1].add_run('Each person enters their share. Running total shows remaining.')
bullet('', bold_prefix='Per-person VC usage: ')
doc.paragraphs[-1].add_run('Each person can choose to apply their personal VCs to their share')
bullet('Table Pot portion distributed: pot split equally before individual VCs applied')

doc.add_heading('8D: Payment Confirmation', level=3)
doc.add_paragraph(
    "In MVP, actual payment happens outside CheckIn (cash, UPI, card via venue's payment system). "
    "CheckIn tracks the credit adjustments."
)
bullet('Summary screen: "Bill settled! \u20b92,800 paid. 200 VCs used. Remaining balance: 1,050 VCs"')
bullet('Quest validation: if bill meets quest thresholds, auto-complete relevant quests')
bullet('"Raise the Stakes" quest check: "You spent \u20b93,000 \u2014 quest complete! +200 VCs earned"')
bullet('Net credit change displayed: "Used 200 VCs, earned 200 VCs = net zero today \U0001f389"')

bold_para('Edge Cases')
bullet('Bill amount disputed: "Contact venue staff to adjust" link')
bullet('Table Pot leftover after bill: remaining pot credits split equally to each diner\'s personal wallet')
bullet('Solo diner (no split needed): Skip split screen, go straight to apply credits \u2192 confirmation')
bullet('POS integration failure: fall back to manual bill entry')

flow_arrow('Bill settled \u2192 Screen 9')

# ──────────────────────────────────────────────────────────────
# SCREEN 9: POST-VISIT / SESSION END
# ──────────────────────────────────────────────────────────────
screen_header('Screen 9: Post-Visit Summary & App Download Prompt',
              'Session wrap-up, achievements, conversion to app')

bold_para('What the diner sees')
bullet('Visit summary card: "Your visit to {Venue}"')
bullet('Credits earned this visit: "+450 VCs earned today"')
bullet('Quests completed: "2 quests completed \u2014 Raise the Stakes, Welcome Aboard"')
bullet('Current balance: "1,250 VCs at {Venue}"')
bullet('Streak tracker: "2-visit streak! Come back Friday for 3x bonus"')

bold_para('App Download CTA (The Conversion Moment)')
doc.add_paragraph(
    "This is the key conversion point. The diner has just had a great experience — credits earned, "
    "quests completed, rewards pending. Now show what they're missing:"
)
bullet('\U0001f513 "See your rank on the leaderboard \u2014 you\'re close to the top!"')
bullet('\U0001f513 "Unlock collectible badges from {Venue}"')
bullet('\U0001f513 "Get personalized restaurant recommendations"')
bullet('\U0001f513 "Discover 50+ venues where your Global Credits work"')
bullet('')
doc.paragraphs[-1].add_run('Big CTA: "Download CheckIn App \u2192"').bold = True
bullet('Secondary: "Maybe later" (dismissible, no penalty)')

bold_para('Post-Visit Notifications (via browser push, if permitted)')
bullet('24h later: "You have 1,250 VCs at {Venue} \u2014 check your rewards!"')
bullet('7 days later: "Your streak expires tomorrow \u2014 visit {Venue} to keep it alive!"')
bullet('30 days before VC expiry: "{X} VCs expiring soon at {Venue}"')

bold_para('Edge Cases')
bullet('Diner closes browser mid-visit: session data persists. On return, resume from last state.')
bullet('Diner doesn\'t settle bill via CheckIn: credits earned from quests still valid. Bill-dependent quests remain incomplete.')
bullet('Push notification permission denied: fall back to email (if provided) or in-app notifications when they open CheckIn next')

doc.add_paragraph('\u2500' * 60)

# ── TECHNICAL FLOW SUMMARY ───────────────────────────────────
doc.add_heading('Technical Flow Summary', level=2)

doc.add_paragraph('End-to-end data flow:')
doc.add_paragraph(
    'QR Scan \u2192 Browser loads venue page (Vercel Edge) \u2192 '
    'Auth via Supabase (phone OTP) \u2192 '
    'GPS validation confirms venue presence \u2192 '
    'Auto-follow: diner linked to venue in Community Graph \u2192 '
    'Active quests loaded (Supabase real-time) \u2192 '
    'Diner completes actions \u2192 '
    'Quest Engine validates completion (POS or manual) \u2192 '
    'Value Engine mints VCs to diner wallet \u2192 '
    'Venue Dashboard updates in real-time \u2192 '
    'Session ends \u2192 App download prompt'
)

doc.add_heading('Key Technical Decisions', level=3)
add_table(
    ['Decision', 'Choice', 'Why'],
    [
        ['Platform', 'Mobile browser (PWA)', 'Zero friction. No app store. QR \u2192 instant access.'],
        ['Auth', 'Phone OTP (Supabase Auth)', 'Universal in India. No Google/social dependency.'],
        ['Real-time', 'Supabase Realtime', 'Table pot, leaderboard, quest progress sync instantly.'],
        ['GPS', 'Browser Geolocation API', 'Anti-fraud. Graceful degradation if denied.'],
        ['Offline', 'Service worker caching', 'Menu and quest data cached. Credits sync when online.'],
        ['POS fallback', 'Manual entry + staff confirm', 'Works without any POS integration for MVP.'],
    ]
)

doc.add_paragraph('\u2500' * 60)

# ── SUCCESS METRICS ──────────────────────────────────────────
doc.add_heading('Success Metrics', level=2)
add_table(
    ['Metric', 'Target', 'Measurement'],
    [
        ['QR \u2192 Landing Page', '>95%', 'Page loads successfully after scan'],
        ['Landing \u2192 Check-in', '>50%', 'Visitors who complete auth + check-in'],
        ['Auth completion time', '<15 seconds', 'Phone entry \u2192 OTP \u2192 checked in'],
        ['Quest engagement', '>60%', 'Checked-in users who view at least 1 quest'],
        ['Quest completion', '>40%', 'Users who complete at least 1 quest per visit'],
        ['VC Redemption', '>60%', 'Earned VCs that get redeemed within 60 days'],
        ['Bill feature usage', '>30%', 'Check-ins that use bill/credit features'],
        ['App download conversion', '>20%', 'Browser users who download the app'],
        ['Return visit (30 day)', '>35%', 'Users who check in again within 30 days'],
        ['Session duration', '>25 min', 'Average time from check-in to session end'],
    ]
)

doc.add_paragraph('\u2500' * 60)

# ── SCREEN MAP ───────────────────────────────────────────────
doc.add_heading('Screen Map (Quick Reference)', level=2)
add_table(
    ['Screen', 'Name', 'Purpose', 'Key Action'],
    [
        ['1', 'Venue Landing', 'First impression, social proof', 'Tap "Check In"'],
        ['2', 'Authentication', 'Phone OTP sign-up/in', 'Enter phone, verify OTP'],
        ['2B', 'Name Entry', 'New user only \u2014 set name', 'Enter first name'],
        ['3', 'Check-In Confirm', 'Celebration + auto-follow', 'Auto-transition (2-3s)'],
        ['4', 'Quest Dashboard', 'Main hub during visit', 'Browse quests, earn credits'],
        ['5', 'Quest Detail', 'Expanded quest view', 'Track progress, complete'],
        ['6', 'Menu', 'Browse menu, quest highlights', 'Find quest-relevant items'],
        ['7', 'Rewards & Wallet', 'VC balance, redeem rewards', 'Redeem VCs for rewards'],
        ['8', 'Bill & Payment', 'Credit usage, bill split', 'Apply credits, split bill'],
        ['9', 'Post-Visit', 'Summary + app download', 'Download app or dismiss'],
    ]
)

# ── Sources ──────────────────────────────────────────────────
doc.add_paragraph('\u2500' * 60)
p = doc.add_paragraph()
p.add_run('Sources: ').italic = True
p.add_run(
    'product/features.md (Phase 1B MVP features), product/overview.md (core loop), '
    'product/architecture.md (data flow + tech stack), product/credit-system.md (VC mechanics).'
).italic = True

# ── Save ─────────────────────────────────────────────────────
output_path = '/Users/ashishnoel/Documents/Projects/Checkin/Checkin_Brain/final-docs/CheckIn — Webapp PRD.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
